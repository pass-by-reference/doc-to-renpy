from docx.document import Document
from docx.text.paragraph import Paragraph

from enum import Enum

from typing import List

class TextType(Enum):
  DIALOGUE = 1
  NARRATION = 2
  SOUND = 3
  NONE = 4

class TextChunk:

  def __init__(self):
    self.paragraphs : List[Paragraph] = []
    self.text_type : TextType = TextType.NONE
    self.character : str = ""

class Consolidate:
  
  def __init__(self, document: Document):

    self.document = document
    self.text_chunks : list[TextChunk] = []
    self.doc_paragraphs : list[Paragraph] = document.paragraphs
  
  def consolidate_paragraphs(self):

    processing_chunk = False

    for paragraph in self.doc_paragraphs:

      if len(paragraph.text) > 0 and paragraph.text[0] == "#":
        continue
      
      if paragraph.text.strip() == "":
        processing_chunk = False
        continue

      if not processing_chunk:
        chunk = TextChunk()
        chunk.paragraphs.append(paragraph)
        chunk.text_type = self.get_text_type(paragraph)
        chunk.character = self.get_character(paragraph, chunk.text_type)

        processing_chunk = True

        self.text_chunks.append(chunk)

        continue

      if processing_chunk:
        self.text_chunks[-1].paragraphs.append(paragraph)

  def get_character(self, paragraph: Paragraph, text_type: TextType) -> str:
    if text_type == TextType.DIALOGUE:
      return paragraph.text.rsplit(":")[0]

    return ""
      
  def get_text_type(self, paragraph: Paragraph) -> TextType:

    if self.is_dialogue(paragraph.text):
      return TextType.DIALOGUE
    elif self.is_sound(paragraph.text):
      return TextType.SOUND
    elif self.is_narration(paragraph.text):
      return TextType.NARRATION
    else:
      return TextType.NONE

  def is_dialogue(self, text: str) -> bool:
    return ":" in text
    
  def is_sound(self, text: str) -> bool:
    return "*" in text

  def is_narration(self, text: str) -> bool:
    return ":" not in text
