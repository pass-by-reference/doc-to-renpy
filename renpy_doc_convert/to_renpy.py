import logging
from collections import namedtuple

from docx.text.run import Run
from docx.shared import RGBColor
from docx.document import Document

from renpy_doc_convert.consolidate import TextChunk, TextType
from typing import List

INDENTATION_SPACES = 2
DEFAULT_FONT_SIZE = 11.0
DEFAULT_FONT_COLOR = "000000" # Hexadecimal Black

# name_found : bool
# text : string
CharNameReturn = namedtuple('CharNameReturn', ['name_found', 'text'])

class ConvertToRenpy:

  def __init__(self, document : Document, chunks: List[TextChunk], output_file_path : str):
    self.chunks : List[TextChunk] = chunks
    self.output_file_path : str = output_file_path
    self.font_standards : FontStandards = FontStandards(document, chunks)
    self.renpy_styler = RenpyStyling(self.font_standards)

    logging.debug("Finish with initializing ConvertToRenpy constructor")

  def get_label(self, output_file_path : str) -> str:
    path_list : List[str] = output_file_path.split('/')
    
    filename_with_extension : str = path_list[len(path_list) - 1]
    label = filename_with_extension.split(".")[0]

    logging.debug("Renpy label: {0}".format(label))

    return label

  def output_renpy_text(self):

    logging.debug("Output renpy text to file")
    with open(self.output_file_path, "w", encoding="utf-8") as file:

      label = self.get_label(self.output_file_path)
      file.write("label {0}:\n\n".format(label))

      logging.debug("Processing {0} text chunk(s)".format(len(self.chunks)))

      for chunk in self.chunks:
        text = self.handle_styling(chunk)
        text = self.handle_escape_characters(text)
        text = self.format_indentation(chunk, text)
        file.write(text)

  def handle_styling(self, chunk: TextChunk) -> str:
    text = ""

    for index, paragraph in enumerate(chunk.paragraphs):

      character_name_found : bool = False
      
      for run in paragraph.runs:

        if(not character_name_found):
          char_ret = self.remove_character_name_in_text(run.text)

          character_name_found : bool = char_ret.name_found
          run.text = char_ret.text

        # Do all bold, italics, and color here
        appendtext = self.renpy_styler.process_run_for_styling(run)
        text = text + appendtext

      if len(chunk.paragraphs) - 1 != index:
        text = text + "\n"
    
    return text

  def handle_escape_characters(self, text : str):
    """
    In renpy, there are special characters that need to be handled

    https://www.renpy.org/doc/html/text.html#escape-characters
    """

    if "\\" in text:
      text = text.replace("\\", "\\\\")

    if "\"" in text:
      text = text.replace('\"', '\\"')

    if "\'" in text:
      text = text.replace("\'", "\\'")

    if "%" in text:
      text = text.replace("%", "\\%")

    return text

  def format_indentation(self, chunk : TextChunk, text: str) -> str:
    if(chunk.text_type == TextType.DIALOGUE):
      return self.format_dialogue(chunk, text)
    else:
      return self.format_non_dialogue(text)

  def format_dialogue(self, chunk : TextChunk, text : str) -> str:

    character = "\"{0}\"".format(chunk.character)
    text = "\"{0}\"".format(text)

    character_text = "{0} {1}".format(character, text)

    # Add spaces to front of text
    indented_text = character_text.rjust(len(character_text) + INDENTATION_SPACES)

    # Add two new lines
    fulltext = "{0}\n\n".format(indented_text)

    return fulltext

  def format_non_dialogue(self, text: str):
    return "  \"" + text + "\"\n\n"

  def remove_character_name_in_text(self, text : str) -> CharNameReturn:
    if ":" in text:
      text = text.split(":", maxsplit=1)[1].lstrip()

      # Returns text without name in it.
      return CharNameReturn(name_found = True, text = text)
    else:
      return CharNameReturn(name_found = False, text = "")

class RenpyStyling:

  def __init__(self, font_stds):
    self.font_stds : FontStandards = font_stds

  def process_run_for_styling(self, run : Run) -> str:
    text = run.text
    if run.bold:
      text = self.convert_bold(text)

    if run.italic:
      text = self.convert_italics(text)

    if run.underline:
      text = self.convert_underline(text)

    if run.font.size:
      text = self.convert_font_size(text, run.font.size.pt)

    if run.font.color and run.font.color.rgb:
      text = self.convert_font_color(text, run.font.color.rgb)

    if run.font.strike:
      text = self.convert_strike(text)
    
    return text
  
  def convert_bold(self, text : str) -> str:
    return "{{b}}{0}{{/b}}".format(text)

  def convert_italics(self, text : str) -> str:
    return "{{i}}{0}{{/i}}".format(text)

  def convert_underline(self, text : str) -> str:
    return "{{u}}{0}{{/u}}".format(text)

  def convert_font_size(self, text : str, run_font_size : int) -> str:
    if self.font_stds.size == run_font_size:
      return text

    diff = run_font_size - self.font_stds.size
    diff = int(diff)

    return "{{size=+{0}}}{1}{{/size}}".format(diff, text)

  def convert_font_color(self, text : str, color : RGBColor):
    if self.font_stds.color == color:
      return text

    rgb_str : str = "#{0}".format(color)

    return "{{color={0}}}{1}{{/color}}".format(rgb_str, text)
    
  def convert_strike(self, text):
    return "{{s}}{0}{{/s}}".format(text)

class FontStandards:
  def __init__(self, document : Document, chunks : TextChunk):
    self.document = document
    self.chunks = chunks
    self.size = self.get_standard_font_size()
    self.color : RGBColor = self.get_standard_font_color()

    logging.debug("Font Size Standard: {0}".format(self.size))
    logging.debug("Font Color standard: {0}".format(self.color))
    logging.debug("Finish with initializing FontStandards constructor")

  def _get_size_first_line(self) -> int:
    if(len(self.chunks) and 
       len(self.chunks[0].paragraphs) and
       len(self.chunks[0].paragraphs[0].runs)
    ):
      run = self.chunks[0].paragraphs[0].runs[0]
      
      if(run.font.size):
        return run.font.size.pt

    return -1

  def _get_document_default(self) -> int:

    if(self.document != None and
       self.document.styles != None and
       self.document.styles.element != None):

      styles_elem = self.document.styles.element

      default_LXML = styles_elem.xpath('w:docDefaults/w:rPrDefault')
      if(len(default_LXML) != 0 and 
        len(default_LXML[0]) != 0):

        # https://github.com/python-openxml/python-docx/blob/master/docx/oxml/text/font.py#L52
        run = default_LXML[0][0] # Should be a type CT_RPr

        if(run.sz_val):
          return run.sz_val.pt

    return -1

  def get_standard_font_size(self) -> int:
    """
    We have to get the standard font size for this document.

    1. Check if the first full line text has a font size. Use that as the priority.
    2. If the first fully line text has no font size, we check the document default
    3. If the document default has no font size, we use hard coded value
    """

    font_size = -1
    
    # Use font from first line
    if(font_size == -1):
      font_size = self._get_size_first_line()

    # If font from first line cannot be found, use document default
    if(font_size == -1):
      font_size = self._get_document_default()

    # If document default is not found, then use hard coded default font
    if(font_size != -1):
      return font_size

    logging.info(
      "Found no font size in document." 
      "Using font size 11.0"
    )
    return DEFAULT_FONT_SIZE 

  def get_standard_font_color(self) -> RGBColor:
    """
    Get standard font color for this document.
    Assume standard font color is black
    """

    return RGBColor.from_string(DEFAULT_FONT_COLOR)